"""
Document parser for handling various document formats (PDF, DOCX, XLSX, etc.).
"""

import logging
from pathlib import Path
from typing import Dict, List, Optional, Union, Any
import asyncio
from concurrent.futures import ThreadPoolExecutor

# Document processing imports
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    import lxml
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

logger = logging.getLogger(__name__)


class DocumentParser:
    """Handles parsing of various document formats."""
    
    def __init__(self, max_workers: int = 4):
        self.max_workers = max_workers
        self.executor = ThreadPoolExecutor(max_workers=max_workers)
        
        # Check available parsers
        self.available_parsers = {
            'pdf': PDF_AVAILABLE,
            'docx': DOCX_AVAILABLE,
            'xlsx': EXCEL_AVAILABLE,
            'html': HTML_AVAILABLE,
            'markdown': MARKDOWN_AVAILABLE
        }
        
        logger.info(f"Available parsers: {[k for k, v in self.available_parsers.items() if v]}")
    
    async def parse_document(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Parse a document and extract its content.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Dictionary containing parsed content and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {"error": "File does not exist", "success": False}
        
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.pdf':
                return await self._parse_pdf_async(file_path)
            elif extension in ['.docx', '.doc']:
                return await self._parse_docx_async(file_path)
            elif extension in ['.xlsx', '.xls']:
                return await self._parse_excel_async(file_path)
            elif extension in ['.html', '.htm']:
                return await self._parse_html_async(file_path)
            elif extension in ['.md', '.markdown']:
                return await self._parse_markdown_async(file_path)
            elif extension in ['.csv']:
                return await self._parse_csv_async(file_path)
            elif extension in ['.json']:
                return await self._parse_json_async(file_path)
            elif extension in ['.xml']:
                return await self._parse_xml_async(file_path)
            else:
                return {
                    "error": f"Unsupported document format: {extension}",
                    "success": False
                }
                
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {e}")
            return {
                "error": str(e),
                "success": False,
                "file_path": str(file_path)
            }
    
    async def _parse_pdf_async(self, file_path: Path) -> Dict[str, Any]:
        """Parse PDF document."""
        if not PDF_AVAILABLE:
            return {"error": "PyPDF2 not available for PDF parsing", "success": False}
        
        def parse_pdf():
            try:
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    
                    content = []
                    metadata = {
                        "num_pages": len(reader.pages),
                        "title": reader.metadata.get('/Title', '') if reader.metadata else '',
                        "author": reader.metadata.get('/Author', '') if reader.metadata else '',
                        "subject": reader.metadata.get('/Subject', '') if reader.metadata else ''
                    }
                    
                    for page_num, page in enumerate(reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text.strip():
                                content.append({
                                    "page": page_num + 1,
                                    "text": page_text
                                })
                        except Exception as e:
                            logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                    
                    full_text = '\n\n'.join([page["text"] for page in content])
                    
                    return {
                        "content": full_text,
                        "pages": content,
                        "metadata": metadata,
                        "format": "pdf",
                        "success": True
                    }
                    
            except Exception as e:
                return {"error": str(e), "success": False}
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, parse_pdf)
    
    async def _parse_docx_async(self, file_path: Path) -> Dict[str, Any]:
        """Parse DOCX document."""
        if not DOCX_AVAILABLE:
            return {"error": "python-docx not available for DOCX parsing", "success": False}
        
        def parse_docx():
            try:
                doc = Document(file_path)
                
                paragraphs = []
                tables = []
                
                # Extract paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append({
                            "text": para.text,
                            "style": para.style.name if para.style else "Normal"
                        })
                
                # Extract tables
                for table_idx, table in enumerate(doc.tables):
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    
                    tables.append({
                        "table_id": table_idx,
                        "data": table_data
                    })
                
                # Combine all text
                full_text = '\n'.join([para["text"] for para in paragraphs])
                
                # Add table content to full text
                for table in tables:
                    table_text = '\n'.join(['\t'.join(row) for row in table["data"]])
                    full_text += f"\n\n[Table {table['table_id']}]\n{table_text}"
                
                return {
                    "content": full_text,
                    "paragraphs": paragraphs,
                    "tables": tables,
                    "format": "docx",
                    "success": True
                }
                
            except Exception as e:
                return {"error": str(e), "success": False}
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, parse_docx)
    
    async def _parse_excel_async(self, file_path: Path) -> Dict[str, Any]:
        """Parse Excel document."""
        if not EXCEL_AVAILABLE:
            return {"error": "openpyxl/pandas not available for Excel parsing", "success": False}
        
        def parse_excel():
            try:
                # Read all sheets
                excel_file = pd.ExcelFile(file_path)
                sheets_data = {}
                
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    # Convert to records
                    records = df.to_dict('records')
                    
                    # Create text representation
                    sheet_text = f"Sheet: {sheet_name}\n"
                    sheet_text += df.to_string(index=False)
                    
                    sheets_data[sheet_name] = {
                        "data": records,
                        "text": sheet_text,
                        "shape": df.shape,
                        "columns": list(df.columns)
                    }
                
                # Combine all sheet text
                full_text = '\n\n'.join([sheet["text"] for sheet in sheets_data.values()])
                
                return {
                    "content": full_text,
                    "sheets": sheets_data,
                    "format": "excel",
                    "success": True
                }
                
            except Exception as e:
                return {"error": str(e), "success": False}
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, parse_excel)
    
    async def _parse_html_async(self, file_path: Path) -> Dict[str, Any]:
        """Parse HTML document."""
        if not HTML_AVAILABLE:
            return {"error": "BeautifulSoup not available for HTML parsing", "success": False}
        
        def parse_html():
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    html_content = file.read()
                
                soup = BeautifulSoup(html_content, 'lxml')
                
                # Extract text content
                text_content = soup.get_text()
                
                # Extract metadata
                title = soup.find('title')
                title_text = title.get_text() if title else ""
                
                meta_tags = soup.find_all('meta')
                metadata = {}
                for meta in meta_tags:
                    name = meta.get('name') or meta.get('property')
                    content = meta.get('content')
                    if name and content:
                        metadata[name] = content
                
                # Extract links
                links = [{"text": a.get_text(), "href": a.get('href')} 
                        for a in soup.find_all('a', href=True)]
                
                # Extract headings
                headings = []
                for i in range(1, 7):
                    for heading in soup.find_all(f'h{i}'):
                        headings.append({
                            "level": i,
                            "text": heading.get_text()
                        })
                
                return {
                    "content": text_content,
                    "title": title_text,
                    "metadata": metadata,
                    "links": links,
                    "headings": headings,
                    "format": "html",
                    "success": True
                }
                
            except Exception as e:
                return {"error": str(e), "success": False}
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, parse_html)
    
    async def _parse_markdown_async(self, file_path: Path) -> Dict[str, Any]:
        """Parse Markdown document."""
        def parse_markdown():
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    md_content = file.read()
                
                # Extract plain text
                plain_text = md_content
                
                # If markdown library is available, convert to HTML first
                if MARKDOWN_AVAILABLE:
                    md = markdown.Markdown(extensions=['meta', 'toc'])
                    html_content = md.convert(md_content)
                    
                    # Extract metadata if available
                    metadata = getattr(md, 'Meta', {})
                    
                    # Parse HTML to get clean text
                    if HTML_AVAILABLE:
                        soup = BeautifulSoup(html_content, 'html.parser')
                        plain_text = soup.get_text()
                else:
                    metadata = {}
                    # Simple markdown parsing
                    import re
                    # Remove markdown syntax for plain text
                    plain_text = re.sub(r'[#*`_~\[\]()]', '', md_content)
                
                # Extract headings
                import re
                headings = []
                for match in re.finditer(r'^(#{1,6})\s+(.+)$', md_content, re.MULTILINE):
                    level = len(match.group(1))
                    text = match.group(2)
                    headings.append({"level": level, "text": text})
                
                return {
                    "content": plain_text,
                    "raw_markdown": md_content,
                    "metadata": metadata,
                    "headings": headings,
                    "format": "markdown",
                    "success": True
                }
                
            except Exception as e:
                return {"error": str(e), "success": False}
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, parse_markdown)
    
    async def _parse_csv_async(self, file_path: Path) -> Dict[str, Any]:
        """Parse CSV file."""
        def parse_csv():
            try:
                if EXCEL_AVAILABLE:
                    df = pd.read_csv(file_path)
                    
                    # Convert to text representation
                    text_content = df.to_string(index=False)
                    
                    return {
                        "content": text_content,
                        "data": df.to_dict('records'),
                        "shape": df.shape,
                        "columns": list(df.columns),
                        "format": "csv",
                        "success": True
                    }
                else:
                    # Fallback to basic CSV parsing
                    import csv
                    with open(file_path, 'r', encoding='utf-8') as file:
                        reader = csv.DictReader(file)
                        data = list(reader)
                        
                        # Create text representation
                        if data:
                            headers = list(data[0].keys())
                            text_content = '\t'.join(headers) + '\n'
                            for row in data:
                                text_content += '\t'.join(str(row[h]) for h in headers) + '\n'
                        else:
                            text_content = ""
                        
                        return {
                            "content": text_content,
                            "data": data,
                            "format": "csv",
                            "success": True
                        }
                
            except Exception as e:
                return {"error": str(e), "success": False}
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, parse_csv)
    
    async def _parse_json_async(self, file_path: Path) -> Dict[str, Any]:
        """Parse JSON file."""
        def parse_json():
            try:
                import json
                with open(file_path, 'r', encoding='utf-8') as file:
                    json_data = json.load(file)
                
                # Convert to readable text
                text_content = json.dumps(json_data, indent=2, ensure_ascii=False)
                
                return {
                    "content": text_content,
                    "data": json_data,
                    "format": "json",
                    "success": True
                }
                
            except Exception as e:
                return {"error": str(e), "success": False}
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, parse_json)
    
    async def _parse_xml_async(self, file_path: Path) -> Dict[str, Any]:
        """Parse XML file."""
        def parse_xml():
            try:
                if HTML_AVAILABLE:
                    with open(file_path, 'r', encoding='utf-8') as file:
                        xml_content = file.read()
                    
                    soup = BeautifulSoup(xml_content, 'xml')
                    text_content = soup.get_text()
                    
                    return {
                        "content": text_content,
                        "raw_xml": xml_content,
                        "format": "xml",
                        "success": True
                    }
                else:
                    # Fallback to basic XML parsing
                    import xml.etree.ElementTree as ET
                    tree = ET.parse(file_path)
                    root = tree.getroot()
                    
                    def extract_text(element):
                        text = element.text or ""
                        for child in element:
                            text += extract_text(child)
                        return text
                    
                    text_content = extract_text(root)
                    
                    return {
                        "content": text_content,
                        "format": "xml",
                        "success": True
                    }
                
            except Exception as e:
                return {"error": str(e), "success": False}
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, parse_xml)
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported document formats."""
        supported = ['.txt', '.csv', '.json', '.xml']  # Always supported
        
        if self.available_parsers['pdf']:
            supported.append('.pdf')
        if self.available_parsers['docx']:
            supported.extend(['.docx', '.doc'])
        if self.available_parsers['xlsx']:
            supported.extend(['.xlsx', '.xls'])
        if self.available_parsers['html']:
            supported.extend(['.html', '.htm'])
        if self.available_parsers['markdown']:
            supported.extend(['.md', '.markdown'])
        
        return supported
    
    def __del__(self):
        """Cleanup executor on deletion."""
        if hasattr(self, 'executor'):
            self.executor.shutdown(wait=False)